import warnings
warnings.simplefilter('ignore')
import numpy as np
import pandas as pd
import os
import re
import sys
import json
import glob
import copy
import math
import string
from io import BytesIO
from collections import Counter
from abc import ABC, abstractmethod
from typing import List, Dict, Union, Any
import faiss
import datrie
from tqdm import tqdm
from docx import Document
from hanziconv import HanziConv
from dataclasses import dataclass
from nltk import word_tokenize
from nltk.stem import PorterStemmer, WordNetLemmatizer
import torch
from transformers import AutoModel, AutoModelForCausalLM, AutoTokenizer, BitsAndBytesConfig
from rag import RagTokenizer

class DocxParser:
    def __init__(self):
        self.rag_tokenizer = RagTokenizer()
    def __extract_table_content(self, tb):
        df = []
        for row in tb.rows:
            df.append([c.text for c in row.cells])
        return self.__compose_table_content(pd.DataFrame(df))
    def __compose_table_content(self, df):
        def blockType(b):
            patt = [
                ("^(20|19)[0-9]{2}[年/-][0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^(20|19)[0-9]{2}年$", "Dt"),
                (r"^(20|19)[0-9]{2}[年/-][0-9]{1,2}月*$", "Dt"),
                ("^[0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^第*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}年*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}[ABCDE]$", "DT"),
                ("^[0-9.,+%/ -]+$", "Nu"),
                (r"^[0-9A-Z/\._~-]+$", "Ca"),
                (r"^[A-Z]*[a-z' -]+$", "En"),
                (r"^[0-9.,+-]+[0-9A-Za-z/$￥%<>（）()' -]+$", "NE"),
                (r"^.{1}$", "Sg")
            ]
            for p, n in patt:
                if re.search(p, b):
                    return n
            try:
                tks = [t for t in self.rag_tokenizer.tokenize(b).split(" ") if len(t) > 1]
                if len(tks) > 3:
                    if len(tks) < 12:
                        return "Tx"
                    else:
                        return "Lx"

                if len(tks) == 1 and self.rag_tokenizer.tag(tks[0]) == "nr":
                    return "Nr"

                return "Ot"
            except:
                return "Ot"
        if len(df) < 2:
            return []
        max_type = Counter([blockType(str(df.iloc[i, j])) for i in range(
            1, len(df)) for j in range(len(df.iloc[i, :]))])
        max_type = max(max_type.items(), key=lambda x: x[1])[0]
        colnm = len(df.iloc[0, :])
        hdrows = [0]  # header is not nessesarily appear in the first line
        if max_type == "Nu":
            for r in range(1, len(df)):
                tys = Counter([blockType(str(df.iloc[r, j]))
                              for j in range(len(df.iloc[r, :]))])
                tys = max(tys.items(), key=lambda x: x[1])[0]
                if tys != max_type:
                    hdrows.append(r)
        lines = []
        for i in range(1, len(df)):
            if i in hdrows:
                continue
            hr = [r - i for r in hdrows]
            hr = [r for r in hr if r < 0]
            t = len(hr) - 1
            while t > 0:
                if hr[t] - hr[t - 1] > 1:
                    hr = hr[t:]
                    break
                t -= 1
            headers = []
            for j in range(len(df.iloc[i, :])):
                t = []
                for h in hr:
                    x = str(df.iloc[i + h, j]).strip()
                    if x in t:
                        continue
                    t.append(x)
                t = ",".join(t)
                if t:
                    t += ": "
                headers.append(t)
            cells = []
            for j in range(len(df.iloc[i, :])):
                if not str(df.iloc[i, j]):
                    continue
                cells.append(headers[j] + str(df.iloc[i, j]))
            lines.append(";".join(cells))
        if colnm > 3:
            return lines
        return ["\n".join(lines)]
    def parse(self, fnm, from_page=0, to_page=100000):
        self.doc = Document(fnm) if isinstance(
            fnm, str) else Document(BytesIO(fnm))
        pn = 0
        secs = []
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            if from_page <= pn < to_page and p.text.strip():
                secs.append((p.text, p.style.name))
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        tbls = [self.__extract_table_content(tb) for tb in self.doc.tables]
        processed_tbl_texts = []
        for block in tbls:
            if not block:
                processed_tbl_texts.append('')
                continue
            text = block[0] if block else ''
            if not text:
                processed_tbl_texts.append('')
                continue
            text = text.replace('    ', ' ').replace('   ', ' ').replace('  ', ' ')
            text = ' '.join(text.splitlines())
            processed_tbl_texts.append(text)
        secs = [sec[0] for sec in secs]
        return secs+processed_tbl_texts